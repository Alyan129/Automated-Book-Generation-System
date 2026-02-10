"""
Export Service for generating output files.
Supports DOCX, PDF, and TXT formats.
"""
from pathlib import Path
from typing import List, Optional
from uuid import UUID
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from src.models.schemas import Book, Chapter
from src.core.config import config
from src.utils.logger import logger


class ExportService:
    """Service for exporting books to various formats"""
    
    def __init__(self):
        """Initialize export service"""
        config.setup_output_dir()
        self.output_dir = config.OUTPUT_DIR
        logger.info(f"Export directory: {self.output_dir}")
    
    def export_to_docx(self, book: Book, chapters: List[Chapter], filename: Optional[str] = None) -> Path:
        """
        Export book to DOCX format.
        
        Args:
            book: Book record
            chapters: List of chapters
            filename: Optional custom filename
            
        Returns:
            Path to exported file
        """
        if not filename:
            filename = f"{book.title.replace(' ', '_')}_{book.id}.docx"
        
        output_path = self.output_dir / filename
        
        try:
            doc = Document()
            
            # Add title page
            title = doc.add_heading(book.title, 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Add metadata
            doc.add_paragraph(f"\nBook ID: {book.id}")
            doc.add_paragraph(f"Generated: {book.created_at.strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_page_break()
            
            # Add chapters
            for chapter in chapters:
                if chapter.content:
                    # Add chapter title
                    doc.add_heading(f"Chapter {chapter.chapter_number}: {chapter.title}", 1)
                    
                    # Add chapter content
                    # Split content by paragraphs and add them
                    paragraphs = chapter.content.split('\n\n')
                    for para in paragraphs:
                        if para.strip():
                            # Skip if it's a duplicate chapter heading
                            if not para.strip().startswith(f"# Chapter {chapter.chapter_number}"):
                                doc.add_paragraph(para.strip())
                    
                    doc.add_page_break()
            
            # Save document
            doc.save(output_path)
            logger.success(f"Exported DOCX: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error exporting to DOCX: {e}")
            raise
    
    def export_to_pdf(self, book: Book, chapters: List[Chapter], filename: Optional[str] = None) -> Path:
        """
        Export book to PDF format.
        
        Args:
            book: Book record
            chapters: List of chapters
            filename: Optional custom filename
            
        Returns:
            Path to exported file
        """
        if not filename:
            filename = f"{book.title.replace(' ', '_')}_{book.id}.pdf"
        
        output_path = self.output_dir / filename
        
        try:
            doc = SimpleDocTemplate(str(output_path), pagesize=letter,
                                   topMargin=0.75*inch, bottomMargin=0.75*inch)
            
            # Container for PDF elements
            story = []
            styles = getSampleStyleSheet()
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                textColor=RGBColor(0, 0, 0),
                spaceAfter=30,
                alignment=TA_CENTER
            )
            
            chapter_style = ParagraphStyle(
                'ChapterTitle',
                parent=styles['Heading1'],
                fontSize=18,
                spaceAfter=12,
                spaceBefore=12
            )
            
            body_style = ParagraphStyle(
                'CustomBody',
                parent=styles['BodyText'],
                fontSize=11,
                alignment=TA_JUSTIFY,
                spaceAfter=12
            )
            
            # Add title page
            story.append(Paragraph(book.title, title_style))
            story.append(Spacer(1, 0.3*inch))
            story.append(Paragraph(f"Book ID: {book.id}", styles['Normal']))
            story.append(Paragraph(f"Generated: {book.created_at.strftime('%Y-%m-%d %H:%M:%S')}", 
                                 styles['Normal']))
            story.append(PageBreak())
            
            # Add chapters
            for chapter in chapters:
                if chapter.content:
                    story.append(Paragraph(
                        f"Chapter {chapter.chapter_number}: {chapter.title}",
                        chapter_style
                    ))
                    
                    paragraphs = chapter.content.split('\n\n')
                    for para in paragraphs:
                        if para.strip() and not para.strip().startswith('#'):
                            # Clean the text for PDF
                            clean_text = para.strip().replace('&', '&amp;')
                            story.append(Paragraph(clean_text, body_style))
                    
                    story.append(PageBreak())
            
            # Build PDF
            doc.build(story)
            logger.success(f"Exported PDF: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error exporting to PDF: {e}")
            raise
    
    def export_to_txt(self, book: Book, chapters: List[Chapter], filename: Optional[str] = None) -> Path:
        """
        Export book to plain text format.
        
        Args:
            book: Book record
            chapters: List of chapters
            filename: Optional custom filename
            
        Returns:
            Path to exported file
        """
        if not filename:
            filename = f"{book.title.replace(' ', '_')}_{book.id}.txt"
        
        output_path = self.output_dir / filename
        
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                # Write title and metadata
                f.write("=" * 80 + "\n")
                f.write(f"{book.title.center(80)}\n")
                f.write("=" * 80 + "\n\n")
                f.write(f"Book ID: {book.id}\n")
                f.write(f"Generated: {book.created_at.strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write("\n" + "=" * 80 + "\n\n")
                
                # Write chapters
                for chapter in chapters:
                    if chapter.content:
                        f.write(f"CHAPTER {chapter.chapter_number}: {chapter.title}\n")
                        f.write("-" * 80 + "\n\n")
                        f.write(chapter.content)
                        f.write("\n\n" + "=" * 80 + "\n\n")
            
            logger.success(f"Exported TXT: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error exporting to TXT: {e}")
            raise
    
    def export_all_formats(self, book: Book, chapters: List[Chapter]) -> dict[str, Path]:
        """
        Export book to all supported formats.
        
        Args:
            book: Book record
            chapters: List of chapters
            
        Returns:
            Dictionary mapping format to file path
        """
        results = {}
        
        try:
            results['docx'] = self.export_to_docx(book, chapters)
        except Exception as e:
            logger.error(f"DOCX export failed: {e}")
        
        try:
            results['pdf'] = self.export_to_pdf(book, chapters)
        except Exception as e:
            logger.error(f"PDF export failed: {e}")
        
        try:
            results['txt'] = self.export_to_txt(book, chapters)
        except Exception as e:
            logger.error(f"TXT export failed: {e}")
        
        return results
